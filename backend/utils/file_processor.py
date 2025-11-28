import fitz  # PyMuPDF
from docx import Document
import openpyxl
import io
from typing import Tuple
import logging

logger = logging.getLogger(__name__)

class FileProcessor:
    """Utility class for processing different file types"""
    
    @staticmethod
    def detect_file_type(file_bytes: bytes, filename: str = "") -> str:
        """
        Определить тип файла по magic bytes и расширению
        """
        # Проверка по magic bytes
        if file_bytes.startswith(b'%PDF'):
            return 'pdf'
        elif file_bytes.startswith(b'PK\x03\x04'):  # ZIP-based formats
            # Дополнительная проверка содержимого
            try:
                # Проверяем наличие характерных папок в ZIP
                if b'word/' in file_bytes[:2000]:
                    return 'docx'
                elif b'xl/' in file_bytes[:2000]:
                    return 'xlsx'
            except Exception:
                pass
        
        # Fallback на расширение файла
        if filename:
            ext = filename.lower().split('.')[-1]
            if ext in ['pdf', 'docx', 'xlsx']:
                return ext
        
        return 'unknown'
    
    @staticmethod
    def process_pdf(file_bytes: bytes) -> Tuple[str, dict]:
        """
        Извлечь текст из PDF файла
        Returns: (extracted_text, metadata)
        """
        doc = None
        try:
            # Открываем PDF из байтов
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            
            text_parts = []
            metadata = {
                "pages": doc.page_count,
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", "")
            }
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():  # Только непустые страницы
                    text_parts.append(f"=== Страница {page_num + 1} ===\n{text}")
            
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"PDF processed: {doc.page_count} pages, {len(full_text)} characters")
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            raise ValueError(f"Failed to process PDF file: {str(e)}")
        finally:
            # Обязательно закрываем документ
            if doc:
                try:
                    doc.close()
                except:
                    pass
    
    @staticmethod
    def process_docx(file_bytes: bytes) -> Tuple[str, dict]:
        """
        Извлечь текст из DOCX файла
        Returns: (extracted_text, metadata)
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            # Основной текст из параграфов
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Текст из таблиц
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Только непустые строки
                        table_data.append(" | ".join(row_data))
                
                if table_data:
                    tables_text.append("=== Таблица ===\n" + "\n".join(table_data))
            
            # Объединяем весь текст
            all_text = []
            if paragraphs:
                all_text.append("=== Основной текст ===\n" + "\n".join(paragraphs))
            if tables_text:
                all_text.extend(tables_text)
            
            full_text = "\n\n".join(all_text)
            
            metadata = {
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            }
            
            logger.info(f"DOCX processed: {len(paragraphs)} paragraphs, {len(doc.tables)} tables")
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            raise ValueError(f"Failed to process DOCX file: {str(e)}")
    
    @staticmethod
    def process_xlsx(file_bytes: bytes) -> Tuple[str, dict]:
        """
        Извлечь данные из Excel файла
        Returns: (extracted_text, metadata)
        """
        try:
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True)
            
            sheets_data = []
            total_rows = 0
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                
                rows_data = []
                sheet_rows = 0
                
                for row in sheet.iter_rows(values_only=True):
                    # Фильтруем пустые ячейки и конвертируем в строки
                    row_values = [str(cell).strip() for cell in row if cell is not None]
                    if row_values:  # Только непустые строки
                        rows_data.append(" | ".join(row_values))
                        sheet_rows += 1
                
                if rows_data:
                    sheet_text = f"=== Лист: {sheet_name} ===\n" + "\n".join(rows_data)
                    sheets_data.append(sheet_text)
                    total_rows += sheet_rows
            
            full_text = "\n\n".join(sheets_data)
            
            metadata = {
                "sheets": len(wb.sheetnames),
                "total_rows": total_rows,
                "sheet_names": wb.sheetnames
            }
            
            wb.close()
            
            logger.info(f"XLSX processed: {len(wb.sheetnames)} sheets, {total_rows} rows")
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"XLSX processing error: {e}")
            raise ValueError(f"Failed to process XLSX file: {str(e)}")
    
    @staticmethod
    def validate_file_size(file_bytes: bytes, max_size: int) -> bool:
        """
        Проверить размер файла
        """
        return len(file_bytes) <= max_size
    
    @staticmethod
    def clean_extracted_text(text: str) -> str:
        """
        Очистить извлеченный текст от лишних символов
        """
        import re
        
        # Убираем лишние пробелы и переносы
        text = re.sub(r'\s+', ' ', text)
        
        # Убираем повторяющиеся символы
        text = re.sub(r'(.)\1{5,}', r'\1\1\1', text)
        
        # Убираем служебные символы
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        return text.strip()
    
    @classmethod
    def process_file(cls, file_bytes: bytes, filename: str) -> Tuple[str, str, dict]:
        """
        Обработать файл любого поддерживаемого типа
        Returns: (extracted_text, file_type, metadata)
        """
        file_type = cls.detect_file_type(file_bytes, filename)
        
        if file_type == 'pdf':
            text, metadata = cls.process_pdf(file_bytes)
        elif file_type == 'docx':
            text, metadata = cls.process_docx(file_bytes)
        elif file_type == 'xlsx':
            text, metadata = cls.process_xlsx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_type}. Supported types: PDF, DOCX, XLSX")
        
        # Очищаем текст
        cleaned_text = cls.clean_extracted_text(text)
        
        return cleaned_text, file_type, metadata